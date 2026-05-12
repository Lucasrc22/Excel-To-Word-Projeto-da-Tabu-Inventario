import io
import os
from datetime import datetime
from typing import Dict, IO, List, Optional, Tuple

import pandas as pd
from django.conf import settings
from django.core.files.storage import default_storage
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate, InlineImage

from .models import Bem


# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------

FIELD_MAP: Dict[str, str] = {
    "":  "imagem_1",
    "A": "imagem_2",
    "B": "imagem_3",
    "C": "imagem_4",
    "D": "imagem_5",
    "E": "imagem_6",
}

IMAGE_LABELS: Dict[str, str] = {
    "":  "Imagem Principal",
    "A": "Imagem A",
    "B": "Imagem B",
    "C": "Imagem C",
    "D": "Imagem D",
    "E": "Imagem E",
}

REQUIRED_COLUMNS: List[str] = [
    "Numero bem",
    "IMOBILIZADO POR AREA",
    "Localizacao",
    "Centro custo",
    "Descricao do CC",
    "Responsavel",
    "Descricao/TAG",
    "Marca",
    "Modelo",
    "Narrativa",
    "Estado Fisico",
]

IMAGE_WIDTH_MM: int = 60
IMAGE_HEIGHT_MM: int = 45


# ---------------------------------------------------------------------------
# Helpers internos — documento Word
# ---------------------------------------------------------------------------

def _set_cell_background(cell, hex_color: str) -> None:
    """Define cor de fundo de uma célula via XML (python-docx não expõe isso nativamente)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_details_table(doc: Document, bem: "Bem") -> None:
    """Tabela de 2 colunas com atributos do bem."""
    data: List[Tuple[str, str]] = [
        ("Número do Bem",   str(bem.numero_bem or "")),
        ("Área",            str(bem.area or "")),
        ("Localização",     str(bem.localizacao or "")),
        ("Centro de Custo", str(bem.centro_custo or "")),
        ("Descrição CC",    str(bem.descricao_cc or "")),
        ("Responsável",     str(bem.responsavel or "")),
        ("Descrição/TAG",   str(bem.descricao_bem or "")),
        ("Marca",           str(bem.marca or "")),
        ("Modelo",          str(bem.modelo or "")),
        ("Narrativa",       str(bem.narrativa or "")),
        ("Estado Físico",   str(bem.estado or "")),
    ]

    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"

    for i, (label, value) in enumerate(data):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]

        label_cell.text = label
        label_cell.paragraphs[0].runs[0].bold = True
        _set_cell_background(label_cell, "D9E1F2")

        value_cell.text = value

    for row in table.rows:
        row.cells[0].width = Mm(55)
        row.cells[1].width = Mm(105)

    doc.add_paragraph()


def _add_images_section(doc: Document, bem: "Bem") -> None:
    """
    Insere bloco fotográfico para todos os sufixos com imagem associada.
    Valida existência e integridade antes de inserir.
    """
    from PIL import Image as PILImage

    imagens_presentes: List[Tuple[str, str]] = [
        (sufixo, campo)
        for sufixo, campo in FIELD_MAP.items()
        if getattr(bem, campo, None) and getattr(bem, campo).name
    ]

    if not imagens_presentes:
        return

    p = doc.add_paragraph()
    run = p.add_run("Registros Fotográficos")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    for sufixo, campo in imagens_presentes:
        field_value = getattr(bem, campo)
        label = IMAGE_LABELS[sufixo]

        try:
            image_path = field_value.path

            if not os.path.exists(image_path):
                doc.add_paragraph(f"⚠ {label}: arquivo não encontrado.")
                continue

            if os.path.getsize(image_path) == 0:
                doc.add_paragraph(f"⚠ {label}: arquivo vazio.")
                continue

            with PILImage.open(image_path) as img:
                img.verify()

            doc.add_picture(image_path, width=Mm(IMAGE_WIDTH_MM), height=Mm(IMAGE_HEIGHT_MM))

        except Exception as exc:
            doc.add_paragraph(
                f"⚠ {label}: erro ao processar — {type(exc).__name__}: {exc}"
            )

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Imagens — upload
# ---------------------------------------------------------------------------

def salvar_imagens(files: List[IO]) -> Dict[str, str]:
    """
    Recebe arquivos de imagem, converte para JPEG e salva via Django storage.
    Retorna { nome_sem_extensao: caminho_relativo }.
    """
    from PIL import Image

    paths: Dict[str, str] = {}
    base_path = "uploads/imagens"

    for f in files:
        try:
            img = Image.open(f)

            if img.mode in ("RGBA", "LA", "P"):
                background = Image.new("RGB", img.size, (255, 255, 255))
                mask = img.split()[-1] if img.mode in ("RGBA", "LA") else None
                background.paste(img, mask=mask)
                img = background
            elif img.mode != "RGB":
                img = img.convert("RGB")

            img_io = io.BytesIO()
            img.save(img_io, format="JPEG", quality=85)
            img_io.seek(0)

            nome = os.path.splitext(f.name)[0]
            file_path = os.path.join(base_path, f"{nome}.jpg")
            saved_path = default_storage.save(file_path, img_io)
            paths[nome] = saved_path

            print(f"[OK] Imagem salva: {saved_path}")

        except Exception as exc:
            print(f"[WARN] Erro ao salvar '{f.name}': {exc}")
            continue

    return paths


# ---------------------------------------------------------------------------
# Importação Excel
# ---------------------------------------------------------------------------

def importar_excel(
    excel_path: str,
    imagens_dict: Dict[str, str],
) -> Tuple[io.BytesIO, str]:
    """
    Lê Excel, persiste Bens no banco e associa imagens para os sufixos
    suportados. Gera documento apenas com os bens da importação atual.
    """
    df = pd.read_excel(excel_path, header=0)
    df.columns = df.columns.str.strip()

    _validar_colunas(df)

    # ⚡ Rastreia apenas os números da planilha atual — evita incluir
    # bens de importações anteriores no documento gerado
    numeros_importados: List[str] = []

    for _, row in df.iterrows():
        numero = str(row["Numero bem"]).strip()
        numeros_importados.append(numero)

        bem, _ = Bem.objects.update_or_create(
            numero_bem=numero,
            defaults={
                "area":          row["IMOBILIZADO POR AREA"],
                "localizacao":   row["Localizacao"],
                "centro_custo":  row["Centro custo"],
                "descricao_cc":  row["Descricao do CC"],
                "responsavel":   row["Responsavel"],
                "descricao_bem": row["Descricao/TAG"],
                "marca":         row["Marca"],
                "modelo":        row["Modelo"],
                "narrativa":     row["Narrativa"],
                "estado":        row["Estado Fisico"],
            },
        )

        _associar_imagens(bem, numero, imagens_dict)

    # Gera documento somente com os bens desta importação
    result = gerar_docx_unificado(numeros_importados)
    deletar_arquivos()
    return result


def _validar_colunas(df: "pd.DataFrame") -> None:
    """Levanta ValueError descritivo se colunas obrigatórias estiverem ausentes."""
    missing = [col for col in REQUIRED_COLUMNS if col not in df.columns]
    if missing:
        raise ValueError(
            f"Colunas ausentes no Excel: {missing}\n"
            f"Colunas disponíveis: {list(df.columns)}"
        )


def _associar_imagens(
    bem: "Bem",
    numero: str,
    imagens_dict: Dict[str, str],
) -> None:
    """
    Associa imagens ao bem para cada sufixo suportado.
    ⚡ save(update_fields=...) atualiza apenas colunas alteradas.
    """
    campos_atualizados: List[str] = []

    for sufixo, campo in FIELD_MAP.items():
        chave = f"{numero}{sufixo}"
        caminho = imagens_dict.get(chave)

        if caminho:
            setattr(bem, campo, caminho)
            campos_atualizados.append(campo)

    if campos_atualizados:
        bem.save(update_fields=campos_atualizados)


# ---------------------------------------------------------------------------
# Geração do documento Word unificado
# ---------------------------------------------------------------------------

def gerar_docx_unificado(
    numeros: Optional[List[str]] = None,
) -> Tuple[io.BytesIO, str]:
    """
    Gera documento Word com os bens da importação atual.

    Args:
        numeros: lista de numero_bem da planilha importada.
                 Se None, inclui todos os bens do banco (retrocompatível).

    Retorna (BytesIO, filename).
    ⚡ Filtra por numero_bem__in para não incluir bens de importações anteriores.
    """
    if numeros:
        bens = Bem.objects.filter(numero_bem__in=numeros).order_by("numero_bem")
    else:
        bens = Bem.objects.all().order_by("numero_bem")

    if not bens.exists():
        raise ValueError("Nenhum bem encontrado para gerar documento.")

    doc = Document()

    # Título principal
    title = doc.add_heading("Inventário de Ativos", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph(
        f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if subtitle.runs:
        subtitle.runs[0].font.size = Pt(10)
        subtitle.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()

    bens_list = list(bens)

    for idx, bem in enumerate(bens_list):

        heading = doc.add_heading(f"Bem  #{bem.numero_bem}", level=1)
        if heading.runs:
            heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        _add_details_table(doc, bem)
        _add_images_section(doc, bem)

        if idx < len(bens_list) - 1:
            doc.add_page_break()

    # Serializa para BytesIO
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    # Persiste cópia em disco com timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"Inventario_Unificado_{timestamp}.docx"
    output_dir = os.path.join(settings.MEDIA_ROOT, "outputs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, output_filename)

    with open(output_path, "wb") as fh:
        fh.write(doc_bytes.getvalue())

    doc_bytes.seek(0)
    return doc_bytes, output_filename


# ---------------------------------------------------------------------------
# Geração individual (compatibilidade legada)
# ---------------------------------------------------------------------------

def gerar_docx(bem: "Bem") -> str:
    """
    Gera documento Word individual via template docxtpl.
    Mantido por compatibilidade — use gerar_docx_unificado() para geração completa.
    """
    template_path = "templates_docx/template.docx"
    doc = DocxTemplate(template_path)

    context: Dict = {
        "numero_bem":    bem.numero_bem,
        "descricao":     bem.descricao_cc,
        "localizacao":   bem.localizacao,
        "centro_custo":  bem.centro_custo,
        "responsavel":   bem.responsavel,
        "marca":         bem.marca,
        "modelo":        bem.modelo,
        "narrativa":     bem.narrativa,
        "estado":        bem.estado,
    }

    for sufixo, campo in FIELD_MAP.items():
        key = f"imagem_{list(FIELD_MAP.keys()).index(sufixo) + 1}"
        field_value = getattr(bem, campo, None)
        context[key] = (
            InlineImage(doc, field_value.path, width=Mm(IMAGE_WIDTH_MM), height=Mm(IMAGE_HEIGHT_MM))
            if field_value else ""
        )

    doc.render(context)

    output_path = f"media/output/{bem.numero_bem}.docx"
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Limpeza de arquivos temporários
# ---------------------------------------------------------------------------

def deletar_arquivos() -> None:
    """Remove todos os arquivos em MEDIA_ROOT/uploads/ após o processamento."""
    upload_dir = os.path.join(settings.MEDIA_ROOT, "uploads")
    for root, _dirs, files in os.walk(upload_dir):
        for file in files:
            try:
                os.remove(os.path.join(root, file))
            except Exception as exc:
                print(f"[WARN] Erro ao deletar '{file}': {exc}")


# ---------------------------------------------------------------------------
# Listagem de documentos gerados
# ---------------------------------------------------------------------------

def listar_documentos_gerados() -> List[Dict[str, str]]:
    """
    Retorna lista de .docx gerados em outputs/, do mais recente ao mais antigo.
    Cada item: { filename, size, date, url_name }
    """
    output_dir = os.path.join(settings.MEDIA_ROOT, "outputs")

    if not os.path.exists(output_dir):
        return []

    files: List[Dict[str, str]] = []

    for filename in sorted(os.listdir(output_dir), reverse=True):
        if not filename.endswith(".docx"):
            continue

        filepath = os.path.join(output_dir, filename)
        file_size = os.path.getsize(filepath)
        file_time = os.path.getmtime(filepath)

        size_str = (
            f"{file_size / (1024 * 1024):.2f} MB"
            if file_size > 1024 * 1024
            else f"{file_size / 1024:.2f} KB"
        )

        files.append(
            {
                "filename": filename,
                "size":     size_str,
                "date":     datetime.fromtimestamp(file_time).strftime("%d/%m/%Y %H:%M:%S"),
                "url_name": filename,
            }
        )

    return files